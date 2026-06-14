from __future__ import annotations

import os
import re
import tempfile
import hashlib
from dataclasses import dataclass
from importlib import import_module
from pathlib import Path
from typing import Callable

import fitz
from fastapi import UploadFile
from docx import Document

from backend.services.pdf_ocr import NormalizedOcrText, extract_pdf_text_with_ocr, normalize_ocr_text


DOCUMENT_DIR = Path(__file__).resolve().parents[2] / "data" / "documents"
MAX_DOCUMENT_SIZE = int(os.getenv("MAX_DOCUMENT_SIZE", str(10 * 1024 * 1024)))
UPLOAD_CHUNK_SIZE = 64 * 1024
PDF_OCR_HARD_MAX_PAGES = int(os.getenv("PDF_OCR_HARD_MAX_PAGES", "100"))
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
WINDOWS_RESERVED_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    "CLOCK$",
    "CONIN$",
    "CONOUT$",
    *(f"COM{number}" for number in range(1, 10)),
    *(f"LPT{number}" for number in range(1, 10)),
}
WINDOWS_UNSAFE_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
SUPERSCRIPT_DIGITS = str.maketrans(
    {"¹": "1", "²": "2", "³": "3", "⁴": "4", "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9"}
)
PARSER_VERSION = "resume-parser-v2"
PROMPT_VERSION = "parse-resume-v2"
SCHEMA_VERSION = "workflow-v2"


@dataclass(frozen=True)
class PageText:
    page_number: int
    text: str
    raw_text: str = ""
    normalized_text: str = ""
    metadata: dict | None = None


@dataclass(frozen=True)
class StoredDocument:
    filename: str
    path: Path
    size: int
    raw_text: str
    pages: list[PageText]


@dataclass(frozen=True)
class StoredUpload:
    filename: str
    path: Path
    size: int


@dataclass(frozen=True)
class DocumentChunk:
    id: str
    run_id: int
    candidate_id: int
    document_type: str
    filename: str
    page_number: int
    section: str
    chunk_index: int
    text: str
    metadata: dict


class UnsupportedDocumentError(ValueError):
    pass


MULTI_POSITION_RE = re.compile(
    r"(职位|岗位|Position|Role)\s*[一二三四五六七八九十\d]+[：:]\s*\S+",
    re.IGNORECASE,
)


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def _safe_filename(filename: str) -> str:
    safe_name = os.path.basename(filename.replace("\\", "/"))
    path = Path(safe_name)
    device_name = safe_name.split(".", 1)[0].translate(SUPERSCRIPT_DIGITS).upper()
    if (
        not safe_name
        or safe_name in {".", ".."}
        or safe_name.startswith(".")
        or safe_name.endswith((" ", "."))
        or path.stem.endswith((" ", "."))
        or WINDOWS_UNSAFE_RE.search(safe_name)
        or device_name in WINDOWS_RESERVED_NAMES
    ):
        raise UnsupportedDocumentError("Invalid document filename")
    if _extension(safe_name) not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(f"Unsupported file type: {safe_name}")
    return safe_name


def _open_exclusive_storage_file(directory: Path, filename: str):
    counter = 1
    candidate = directory / filename
    while True:
        try:
            flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL
            if hasattr(os, "O_BINARY"):
                flags |= os.O_BINARY
            fd = os.open(candidate, flags, 0o600)
            return candidate, candidate.name, os.fdopen(fd, "wb")
        except FileExistsError:
            stored_name = f"{Path(filename).stem}_{counter}{Path(filename).suffix}"
            candidate = directory / stored_name
            counter += 1


async def _write_upload_chunks(file: UploadFile, stream, filename: str) -> int:
    total = 0
    read_size = min(UPLOAD_CHUNK_SIZE, MAX_DOCUMENT_SIZE + 1)
    while chunk := await file.read(read_size):
        total += len(chunk)
        if total > MAX_DOCUMENT_SIZE:
            raise UnsupportedDocumentError(
                f"{filename} exceeds maximum size of {MAX_DOCUMENT_SIZE} bytes"
            )
        stream.write(chunk)
    if not total:
        raise UnsupportedDocumentError(f"{filename} is empty")
    return total


async def store_and_extract_upload(
    file: UploadFile,
    storage_root: str | Path = DOCUMENT_DIR,
    progress_callback: Callable[[int, int], None] | None = None,
) -> StoredDocument:
    stored = await store_upload(file, storage_root=storage_root)
    try:
        pages = extract_stored_pages(stored.path, stored.filename, progress_callback=progress_callback)
        raw_text = "\n".join(page.text for page in pages).strip()
        _ensure_text(raw_text, stored.filename)
        return StoredDocument(
            filename=stored.filename,
            path=stored.path,
            size=stored.size,
            raw_text=raw_text,
            pages=pages,
        )
    except Exception:
        _best_effort_delete(stored.path)
        raise


async def store_upload(
    file: UploadFile,
    storage_root: str | Path = DOCUMENT_DIR,
) -> StoredUpload:
    filename = _safe_filename(file.filename or "")
    directory = Path(storage_root)
    directory.mkdir(parents=True, exist_ok=True)
    path: Path | None = None
    try:
        path, stored_filename, stream = _open_exclusive_storage_file(directory, filename)
        with stream:
            size = await _write_upload_chunks(file, stream, stored_filename)
        return StoredUpload(
            filename=stored_filename,
            path=path,
            size=size,
        )
    except Exception:
        _best_effort_delete(path)
        raise


def extract_stored_pages(
    path: str | Path,
    filename: str,
    progress_callback: Callable[[int, int], None] | None = None,
) -> list[PageText]:
    document_path = Path(path)
    safe_filename = _safe_filename(filename)
    ext = _extension(safe_filename)
    try:
        if ext == ".pdf":
            return _extract_pdf(document_path, safe_filename, progress_callback=progress_callback)
        if ext == ".docx":
            text = _extract_docx(document_path)
        elif ext == ".doc":
            return _extract_legacy_doc(document_path, safe_filename)
        else:
            text = _extract_text_file(document_path, safe_filename)
    except UnsupportedDocumentError:
        raise
    except Exception as exc:
        raise UnsupportedDocumentError(f"Failed to extract text from {safe_filename}: {exc}") from exc

    _ensure_text(text, safe_filename)
    if progress_callback:
        progress_callback(1, 1)
    return [PageText(page_number=1, text=text)]


def delete_stored_file(path: str | Path | None) -> None:
    if not path:
        return
    Path(path).unlink(missing_ok=True)


def _best_effort_delete(path: str | Path | None) -> None:
    try:
        delete_stored_file(path)
    except OSError:
        pass


async def extract_upload_pages(file: UploadFile) -> list[PageText]:
    filename = _safe_filename(file.filename or "")
    ext = _extension(filename)
    tmp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp_path = Path(tmp.name)
            await _write_upload_chunks(file, tmp, filename)
        return extract_stored_pages(tmp_path, filename)
    finally:
        _best_effort_delete(tmp_path)


def _extract_pdf(
    path: Path,
    filename: str,
    progress_callback: Callable[[int, int], None] | None = None,
) -> list[PageText]:
    doc = fitz.open(path)
    native_pages: dict[int, str] = {}
    visual_pages: set[int] = set()
    page_count = len(doc)
    try:
        for index, page in enumerate(doc, start=1):
            native_pages[index] = (page.get_text("text") or "").strip()
            if page.get_images(full=True) or page.get_drawings():
                visual_pages.add(index)
            if progress_callback:
                progress_callback(index, page_count)
    finally:
        doc.close()

    ocr_candidates = {
        page_number
        for page_number, text in native_pages.items()
        if _needs_page_ocr(text, has_visual_content=page_number in visual_pages)
    }
    if not ocr_candidates:
        pages = [PageText(page_number=number, text=text) for number, text in native_pages.items()]
        combined = "\n".join(page.text for page in pages)
        _ensure_text(combined, filename)
        return pages

    last_missing_page = max(ocr_candidates)
    if last_missing_page > PDF_OCR_HARD_MAX_PAGES:
        raise UnsupportedDocumentError(
            f"{filename} requires OCR through page {last_missing_page}, "
            f"which exceeds the hard limit of {PDF_OCR_HARD_MAX_PAGES} pages"
        )
    try:
        ocr_result = extract_pdf_text_with_ocr(path, max_pages=last_missing_page)
    except Exception as exc:
        raise UnsupportedDocumentError(f"{filename} OCR is unavailable or failed: {exc}") from exc
    ocr_pages = {page.page_number: page.text for page in _parse_ocr_pages(ocr_result.text)}
    required_ocr_pages = {
        page_number
        for page_number in ocr_candidates
        if not _has_usable_native_text(native_pages[page_number])
    }
    if not ocr_pages and len(required_ocr_pages) == page_count:
        raise UnsupportedDocumentError(f"{filename} OCR produced no extractable text")
    uncovered = sorted(page_number for page_number in required_ocr_pages if page_number not in ocr_pages)
    if uncovered:
        missing = ", ".join(str(page_number) for page_number in uncovered)
        raise UnsupportedDocumentError(f"{filename} OCR did not cover missing page {missing}")
    for page_number in required_ocr_pages:
        _ensure_text(ocr_pages[page_number], filename)

    pages = [
        PageText(
            page_number=page_number,
            text=_merge_page_text(native_pages[page_number], ocr_pages.get(page_number, ""))
            if page_number in ocr_candidates
            else native_pages[page_number],
        )
        for page_number in range(1, page_count + 1)
    ]
    _ensure_text("\n".join(page.text for page in pages), filename)
    return pages


def _needs_page_ocr(text: str, *, has_visual_content: bool) -> bool:
    stripped = (text or "").strip()
    min_chars = int(os.getenv("PDF_TEXT_MIN_CHARS", "30"))
    return has_visual_content and len(stripped) < min_chars


def _has_usable_native_text(text: str) -> bool:
    return bool(re.search(r"[\w\u4e00-\u9fff]", (text or "").strip()))


def _merge_page_text(native_text: str, ocr_text: str) -> str:
    merged: list[str] = []
    seen: set[str] = set()
    for line in [*native_text.splitlines(), *ocr_text.splitlines()]:
        clean_line = line.strip()
        normalized = re.sub(r"\s+", " ", clean_line).casefold()
        if clean_line and normalized not in seen:
            merged.append(clean_line)
            seen.add(normalized)
    return "\n".join(merged)


OCR_PAGE_RE = re.compile(
    r"^--- OCR Page (\d+) ---\s*\n(.*?)(?=^--- OCR Page \d+ ---|\Z)",
    re.MULTILINE | re.DOTALL,
)


def _parse_ocr_pages(text: str) -> list[PageText]:
    return [
        PageText(page_number=int(match.group(1)), text=match.group(2).strip())
        for match in OCR_PAGE_RE.finditer(text or "")
        if match.group(2).strip()
    ]


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    return "\n".join([*paragraphs, *table_lines]).strip()


def _extract_text_file(path: Path, filename: str) -> str:
    payload = path.read_bytes()
    decode_errors: list[str] = []
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return payload.decode(encoding).strip()
        except UnicodeDecodeError as exc:
            decode_errors.append(f"{encoding}: {exc}")
    raise UnsupportedDocumentError(f"Failed to decode text from {filename}: {'; '.join(decode_errors)}")


def partition_doc(*, filename: str):
    parser = import_module("unstructured.partition.doc").partition_doc
    return parser(filename=filename)


def _extract_legacy_doc(path: Path, filename: str) -> list[PageText]:
    try:
        elements = partition_doc(filename=str(path))
    except Exception as exc:
        raise UnsupportedDocumentError(
            f"Failed to extract legacy Word document {filename}; "
            f"the .doc parser or its external dependencies may be unavailable: {exc}"
        ) from exc

    text = "\n".join(str(element).strip() for element in elements if str(element).strip())
    _ensure_text(text, filename)
    return [PageText(page_number=1, text=text)]


def _ensure_text(text: str, filename: str) -> None:
    if len(text.strip()) < 20:
        raise UnsupportedDocumentError(f"{filename} has no extractable text")
    assert_acceptable_text_quality(text, filename=filename)


def assert_acceptable_text_quality(text: str, *, filename: str = "document") -> None:
    stripped = (text or "").strip()
    if len(stripped) < 20:
        raise UnsupportedDocumentError(f"TEXT_QUALITY_TOO_LOW: {filename} has too little readable text")
    printable = sum(1 for char in stripped if char.isprintable() and not char.isspace())
    non_space = sum(1 for char in stripped if not char.isspace())
    printable_ratio = printable / max(non_space, 1)
    replacement_ratio = stripped.count("\ufffd") / max(len(stripped), 1)
    mojibake_tokens = ("锟", "斤拷", "ï¿½", "Ã", "Â")
    mojibake_count = sum(stripped.count(token) for token in mojibake_tokens)
    mojibake_ratio = mojibake_count / max(len(stripped), 1)
    cjk_count = sum(1 for char in stripped if "\u4e00" <= char <= "\u9fff")
    ascii_alpha = sum(1 for char in stripped if char.isascii() and char.isalpha())
    symbol_count = sum(1 for char in stripped if not char.isalnum() and not char.isspace())
    symbol_ratio = symbol_count / max(len(stripped), 1)
    long_token = any(len(token) > 120 for token in re.split(r"\s+", stripped))
    if (
        printable_ratio < 0.85
        or replacement_ratio > 0.01
        or mojibake_ratio > 0.02
        or symbol_ratio > 0.45
        or long_token
        or (len(stripped) > 120 and cjk_count == 0 and ascii_alpha < 20)
    ):
        raise UnsupportedDocumentError(f"TEXT_QUALITY_TOO_LOW: {filename} appears unreadable")


def detect_multiple_positions(text: str) -> bool:
    normalized = text or ""
    if len(MULTI_POSITION_RE.findall(normalized)) >= 2:
        return True
    title_markers = re.findall(r"(?im)^\s*(职位|岗位|Position|Role)\s*[:：]", normalized)
    responsibility_markers = re.findall(r"(?im)^\s*(岗位职责|职责|Responsibilities)\s*[:：]", normalized)
    return len(title_markers) >= 2 or len(responsibility_markers) >= 2


SECTION_ALIASES = {
    "教育经历": "教育经历",
    "教育背景": "教育经历",
    "实习经历": "实习经历",
    "工作经历": "工作经历",
    "科研经历": "科研经历",
    "项目经历": "项目经历",
    "项目经验": "项目经历",
    "校园经历": "校园经历",
    "专业技能": "专业技能",
    "技能证书": "技能证书",
    "荣誉奖项": "荣誉奖项",
    "自我评价": "自我评价",
    "个人总结": "个人总结",
    "education": "Education",
    "experience": "Experience",
    "workexperience": "Experience",
    "internshipexperience": "Experience",
    "projects": "Projects",
    "projectexperience": "Projects",
    "skills": "Skills",
    "professionalskills": "Skills",
    "awards": "Awards",
    "honors": "Awards",
    "summary": "Summary",
}
WORK_SECTIONS = {"实习经历", "工作经历", "Experience"}
BULLET_RE = re.compile(r"^\s*[-*•·]\s*(.+)$")
DATE_RANGE_RE = re.compile(
    r"((?:20\d{2}|19\d{2})[./-]\d{1,2}|(?:20\d{2}|19\d{2}))[至~\-—–到 ]+((?:20\d{2}|19\d{2})[./-]\d{1,2}|(?:20\d{2}|19\d{2})|至今|Present)",
    re.IGNORECASE,
)


def _canonical_section(line: str) -> str | None:
    stripped = line.strip().strip("#").strip().strip("：:")
    if not stripped or len(stripped) > 40:
        return None
    compact = re.sub(r"[\s:：|/\\_-]+", "", stripped).casefold()
    return SECTION_ALIASES.get(compact) or SECTION_ALIASES.get(stripped)


def chunk_pages(
    *,
    pages: list[PageText],
    run_id: int,
    candidate_id: int,
    document_type: str,
    filename: str,
    chunk_size: int = 500,
    overlap: int = 50,
) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    chunk_index = 0
    current_section = "正文"
    section_lines: list[tuple[str, int, str, str, list[dict]]] = []
    work_context: dict[str, str] = {}
    pending_work_lines: list[tuple[str, int, str, str, list[dict]]] = []

    def append_chunk_from_lines(
        lines: list[tuple[str, int, str, str, list[dict]]],
        section: str,
        *,
        entity_type: str = "",
        entity_id: str = "",
        prefix_context: dict[str, str] | None = None,
    ) -> None:
        # Work-experience bullets often lose employer/title/date context after
        # OCR and chunking. Prefixing recovered context keeps each vector chunk
        # independently useful for later evidence retrieval.
        nonlocal chunk_index
        if not lines:
            return
        text = "\n".join(item[0] for item in lines).strip()
        if prefix_context is not None:
            text = _with_work_prefix(text, section, prefix_context)
        raw_text = "\n".join(item[2] for item in lines).strip()
        normalized_text = "\n".join(item[3] for item in lines).strip()
        ambiguities = [entry for item in lines for entry in item[4]]
        page_start = min(item[1] for item in lines)
        page_end = max(item[1] for item in lines)
        for part in _split_text_units(text, chunk_size):
            chunk_index = _append_chunk(
                chunks,
                run_id,
                candidate_id,
                document_type,
                filename,
                page_start,
                section,
                chunk_index,
                part,
                page_end=page_end,
                raw_text=raw_text,
                normalized_text=normalized_text,
                entity_type=entity_type,
                entity_id=entity_id,
                extraction_method="ocr",
                ambiguities=ambiguities,
            )

    def flush_section_lines() -> None:
        nonlocal section_lines
        if not section_lines:
            return
        append_chunk_from_lines(section_lines, current_section)
        section_lines = []

    def flush_pending_work() -> None:
        nonlocal pending_work_lines
        if not pending_work_lines:
            return
        append_chunk_from_lines(
            pending_work_lines,
            current_section,
            entity_type="work",
            entity_id=_work_entity_id(work_context),
            prefix_context=work_context,
        )
        pending_work_lines = []

    for page in pages:
        normalized_page = _normalized_page_text(page)
        lines = [line.strip() for line in normalized_page.normalized_text.splitlines() if line.strip()]
        for line in lines:
            next_section = _canonical_section(line)
            if next_section:
                flush_pending_work()
                flush_section_lines()
                current_section = next_section
                if current_section in WORK_SECTIONS:
                    work_context = {}
                continue

            if current_section in WORK_SECTIONS:
                bullet = _bullet_text(line)
                line_tuple = (
                    bullet or line,
                    page.page_number,
                    bullet or line,
                    bullet or line,
                    normalized_page.ambiguities,
                )
                if _looks_like_work_header(line):
                    flush_pending_work()
                    work_context = {}
                    _update_work_context(work_context, line)
                    continue
                if bullet:
                    flush_pending_work()
                    pending_work_lines.append(line_tuple)
                    flush_pending_work()
                    continue
                if _work_context_complete(work_context):
                    pending_work_lines.append(line_tuple)
                    continue
                _update_work_context(work_context, line)
                if _work_context_complete(work_context):
                    continue
                if _has_work_context(work_context):
                    pending_work_lines.append(line_tuple)
                    continue

            section_lines.append(
                (
                    line,
                    page.page_number,
                    line,
                    line,
                    normalized_page.ambiguities,
                )
            )
    flush_pending_work()
    flush_section_lines()
    return chunks


def _normalized_page_text(page: PageText) -> NormalizedOcrText:
    if page.normalized_text:
        return NormalizedOcrText(
            raw_text=page.raw_text or page.text,
            normalized_text=page.normalized_text,
            ambiguities=(_page_metadata(page).get("ocr_ambiguities") or []),
        )
    return normalize_ocr_text(page.text)


def _page_metadata(page: PageText) -> dict:
    return page.metadata or {}


def _page_extraction_method(_ambiguities: list[dict]) -> str:
    return "ocr"


def _bullet_text(line: str) -> str:
    match = BULLET_RE.match(line)
    return match.group(1).strip() if match else ""


def _update_work_context(context: dict[str, str], line: str) -> None:
    date_match = DATE_RANGE_RE.search(line)
    if date_match:
        context["time"] = date_match.group(0)
        line = DATE_RANGE_RE.sub("", line).strip(" -｜|")
    if line and "company" not in context:
        parts = re.split(r"\s{2,}|[｜|]", line)
        if len(parts) >= 2:
            context["company"] = parts[0].strip()
            context["title"] = parts[1].strip()
        else:
            context["company"] = line
    elif line and "title" not in context:
        context["title"] = line


def _looks_like_work_header(line: str) -> bool:
    return bool(DATE_RANGE_RE.search(line) and re.search(r"[|丨｜]", line))


def _work_context_complete(context: dict[str, str]) -> bool:
    return bool(context.get("company") and context.get("title") and context.get("time"))


def _has_work_context(context: dict[str, str]) -> bool:
    return bool(context.get("company") or context.get("title") or context.get("time"))


def _with_work_prefix(text: str, section: str, context: dict[str, str]) -> str:
    prefix = [
        f"章节：{section}",
        f"公司：{context.get('company', '')}",
        f"职位：{context.get('title', '')}",
        f"时间：{context.get('time', '')}",
    ]
    return "\n".join([*prefix, text]).strip()


def _work_entity_id(context: dict[str, str]) -> str:
    return "|".join(
        part for part in [context.get("company", ""), context.get("title", ""), context.get("time", "")] if part
    )


def _split_text_units(text: str, chunk_size: int) -> list[str]:
    if len(text) <= chunk_size:
        return [text]
    sentences = [item.strip() for item in re.split(r"(?<=[。！？.!?])\s*", text) if item.strip()]
    if len(sentences) <= 1:
        return [text]
    parts: list[str] = []
    buffer = ""
    for sentence in sentences:
        if buffer and len(buffer) + len(sentence) + 1 > chunk_size:
            parts.append(buffer.strip())
            buffer = sentence
        else:
            buffer = f"{buffer}\n{sentence}".strip()
    if buffer:
        parts.append(buffer.strip())
    return parts


def _append_chunk(
    chunks: list[DocumentChunk],
    run_id: int,
    candidate_id: int,
    document_type: str,
    filename: str,
    page_number: int,
    section: str,
    chunk_index: int,
    text: str,
    *,
    page_end: int | None = None,
    raw_text: str = "",
    normalized_text: str = "",
    entity_type: str = "",
    entity_id: str = "",
    extraction_method: str = "native",
    ambiguities: list[dict] | None = None,
) -> int:
    clean_text = text.strip()
    if not clean_text:
        return chunk_index
    normalized_value = normalized_text or clean_text
    embedding_text = "\n".join(
        part
        for part in [
            f"section: {section or 'body'}",
            f"entity_type: {entity_type}" if entity_type else "",
            f"entity_id: {entity_id}" if entity_id else "",
            normalized_value,
        ]
        if part
    )
    chunk_id = f"{run_id}:{candidate_id or 'jd'}:{document_type}:{chunk_index}"
    chunks.append(
        DocumentChunk(
            id=chunk_id,
            run_id=run_id,
            candidate_id=candidate_id,
            document_type=document_type,
            filename=filename,
            page_number=page_number,
            section=section or "正文",
            chunk_index=chunk_index,
            text=clean_text,
            metadata={
                "section": section or "正文",
                "entity_type": entity_type,
                "entity_id": entity_id,
                "page_start": page_number,
                "page_end": page_end or page_number,
                "extraction_method": extraction_method,
                "normalized": True,
                "text_length": len(clean_text),
                "raw_text": raw_text or clean_text,
                "normalized_text": normalized_value,
                "embedding_text": embedding_text,
                "ocr_ambiguities": ambiguities or [],
                "parser_version": PARSER_VERSION,
                "prompt_version": PROMPT_VERSION,
                "schema_version": SCHEMA_VERSION,
                "ocr_engine": "rapidocr",
                "ocr_version": os.getenv("OCR_VERSION", "unknown") or "unknown",
                "embedding_model": os.getenv("EMBEDDING_MODEL", "unknown") or "unknown",
                "embedding_version": os.getenv("EMBEDDING_VERSION", "unknown") or "unknown",
                "normalized_text_hash": hashlib.sha256(normalized_value.encode("utf-8")).hexdigest(),
            },
        )
    )
    return chunk_index + 1
